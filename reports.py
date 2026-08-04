"""Report generation: data assembly + PDF/Excel rendering.

"Report" is one of PI3 Plant Edition's own standard, always-included
capabilities (see pages/10_PI3_AI_Connectivity.py's docstring: "Standard
version (always included): Search, Compare, Retrieve, Structure, Report,
Review and Approval.") - this module is what had been missing. It is not
gated behind PI3 connectivity; every logged-in user can generate reports.

Two report types that predate the 2026-08-04 Reports redesign, each with a
data-assembly function (plain dict, no Streamlit import, easy to unit
test) and a PDF + Excel renderer pair:

- build_period_summary_data() / render_period_summary_pdf() / render_period_summary_excel()
  One plant/product family/date range: KPIs, pass rate, recurring issues,
  the run list, and a breakdown by foam grade.
- build_trial_report_data() / render_trial_report_pdf() / render_trial_report_excel()
  One closed Customer Trial or Optimization Trial (see db.CustomerTrial /
  db.OptimizationTrial - the two independent lab-trial flows, added
  2026-08-03): objective/hypothesis, what changed, outcome/conclusion,
  reviewer sign-off - a formal closeout writeup. Rebuilt 2026-08-04 to
  cover these two (the trials people actually use) after the old
  TrialRecord concept (a formal-experiment flag on a production run) was
  removed - zero real rows across 244 production runs, fully superseded
  by these two independent, self-contained closeout flows.

pages/21_Report.py wires these to selectors, an in-app preview, and
st.download_button for both file formats.

Two purpose-built reports for the Recipes page (pages/3_Recipe_Version_
Record.py), added 2026-08-04 as the first output of the app-wide Reports
redesign (see PI3_Gaps note: reports must be aggregated/purpose-built
answers to a specific question, never a raw-data dump or a PI3-narrative
document - the CSV export on every table already covers raw-data needs,
and PI3's own Word download on relevant pages already covers narrative):

- build_recipe_formulation_record_data() / render_recipe_formulation_record_pdf()
  / render_recipe_formulation_record_excel()
  One recipe version: the formulation itself (materials/php/supplier/
  role), its quality specs vs. actual results aggregated over a chosen
  date range (across every production run/customer trial/optimization
  trial built on this recipe version), and cost per kg - "is this
  formulation meeting spec, and what does it cost" in one document for
  internal use/approval (not customer-facing - a customer-facing version
  would need to omit the formulation itself).
- build_where_used_report_data() / render_where_used_report_pdf()
  / render_where_used_report_excel()
  One raw material: every recipe version (active and retired) that uses
  it with its php/role, the target properties of every foam grade
  affected, and any Customer/Optimization Trial precedent tied to a
  recipe version containing it - "if I replace this material, what's
  affected and what trials already exist to lean on."

A third purpose-built report, added 2026-08-04 for one production run -
REPLACING the earlier build_run_report_data() / render_run_report_pdf() /
render_run_report_excel() (removed the same day: a flat header plus four
raw tables - recipe components, process settings, quality results,
quality issues - with no synthesis, exactly the "factual but not adding
value" pattern the whole Reports redesign exists to fix). Lives on the
Report page (pages/21_Report.py), not the Production Run page itself: per
user direction, a report whose subject is a single simple choice (pick
one run from a dropdown) belongs on the Report page alongside the other
selector-driven reports; a report whose subject needs a comprehensive
multi-field selection first (date range, foam grade, etc. - see the
Quality Test Result report below) belongs on its own page instead, next
to where that selection naturally happens.

- build_batch_release_record_data() / render_batch_release_record_pdf()
  / render_batch_release_record_excel()
  One production run: the recipe used (in full - not just a reference),
  and a rolled-up quality conformance verdict (Pass/Fail per tested
  property plus one overall Conforming/Non-conforming/Incomplete verdict)
  and any quality issues recorded. If - and only if - a flag is raised
  (a failed result or a recorded quality issue), the report widens to
  pull supporting context from every other tab on the Production Run
  page: Setup-vs-Finalized process-setting deviations (including
  fall-plate position changes), the Finalized phase's actual component
  stream readings (with any non-"Valid" calibration status called out),
  and any Production Events logged during the run - "does this batch
  look wrong, and if so what else was going on at the time" in one
  document, not five separate tab exports. A clean run stays a short
  document; a flagged one pulls in exactly what's relevant, not
  everything that exists.

A fifth report, added 2026-08-04, lives on its own page rather than the
Report page - per the same placement principle stated above, this report's
subject is a comprehensive multi-field selection (Pass/Fail, Property, and
Foam scope), not a single dropdown choice, so it lives on
pages/5_Physical_Property_Result.py, right below the filter controls and
the existing on-page Pareto chart it shares its scope with:

- build_quality_test_report_data() / render_quality_test_report_pdf()
  / render_quality_test_report_excel()
  Takes the exact set of PhysicalPropertyResult rows the page has already
  scoped (tenant) and filtered (Pass/Fail, Property, Foam scope) - the
  report never re-derives its own selection, so it always matches what's
  on screen. Aggregates that set into a pass-rate summary, a failures-by-
  property breakdown (bar chart), a pass/fail-by-foam-grade breakdown
  (bar chart, only shown when the selection spans more than one grade),
  and a curated table of just the failing results (target/actual/
  deviation) - not the full underlying row set, which the page's own CSV
  export already covers.

A sixth, narrower report type lives here too:

- build_pi3_qa_report_data() / render_pi3_qa_report_docx()
  A single "Ask PI3" question-and-answer exchange (see
  helpers.render_ask_pi3_section) - the question, PI3's answer, and an
  appendix of the exact data PI3 checked to produce it (SQL + rows
  returned, or the verified-analysis arguments and result). Unlike the
  reports above, this is DOCX only (no PDF/Excel - there's no
  tabular data here that benefits from a spreadsheet), and it is always
  built from this same code path, so every export has identical
  formatting regardless of who generates it or what was asked - a
  hand-maintained Word template would drift over time; this can't.
"""

import datetime as dt
import io
import os
import re

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from openpyxl.chart import BarChart, Reference
from openpyxl.utils import get_column_letter
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.shapes import Drawing
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from analytics import PHASE_SETTING_FIELDS, PHASE_SETTING_LABELS, recipe_version_cost
from db import (
    ComponentStreamReading,
    CustomerTrial,
    FallplateSectionPosition,
    FoamGrade,
    FoamGradeTargetProperty,
    Machine,
    OptimizationTrial,
    Plant,
    PhysicalPropertyResult,
    ProductFamily,
    ProductionEvent,
    ProductionPhase,
    ProductionRun,
    QualityObservation,
    RawMaterial,
    RecipeComponent,
    RecipeVersion,
)
from quality_standards import compute_pass_fail

STYLES = getSampleStyleSheet()


# ---------------------------------------------------------------------------
# Shared rendering helpers
# ---------------------------------------------------------------------------

def _pdf_bytes(build_story):
    """build_story(story: list) appends reportlab flowables to story."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
        leftMargin=16 * mm, rightMargin=16 * mm,
    )
    story = []
    build_story(story)
    doc.build(story)
    return buf.getvalue()


def _excel_bytes(sheets, charts=None):
    """sheets: dict of sheet_name -> list-of-dicts or DataFrame. Empty
    sections still get a sheet (with a placeholder row) so the workbook
    structure is predictable regardless of what data exists.

    charts: optional list of (sheet_name, chart_builder) pairs. Each
    chart_builder(ws, df) is called after every sheet has been written -
    ExcelWriter needs the full workbook before a chart can reference
    another sheet's cells - with the openpyxl worksheet and the exact
    DataFrame written to it, so it can compute correct cell ranges. See
    _add_bar_chart for the standard builder."""
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        written = {}
        for name, rows in sheets.items():
            df = rows if isinstance(rows, pd.DataFrame) else pd.DataFrame(rows)
            if df.empty:
                df = pd.DataFrame({"—": ["No data recorded"]})
            # Excel sheet names are capped at 31 characters.
            sheet_name = name[:31]
            df.to_excel(writer, sheet_name=sheet_name, index=False)
            written[name] = (writer.sheets[sheet_name], df)
        for name, chart_builder in (charts or []):
            ws_df = written.get(name)
            if ws_df:
                chart_builder(*ws_df)
    return buf.getvalue()


def _add_bar_chart(ws, df, title, category_col, value_cols):
    """Adds a native Excel bar chart to worksheet ws, built from the data
    just written to it - category_col plus one or more value_cols (each
    becomes its own series, e.g. Pass count / Fail count side by side).
    Anchored a couple of columns to the right of the data table so it
    doesn't overlap it. No-ops on the placeholder "No data recorded" sheet
    (category_col won't be a real column in that case)."""
    if df.empty or category_col not in df.columns:
        return
    n_rows = len(df)
    col_idx = {c: i + 1 for i, c in enumerate(df.columns)}
    chart = BarChart()
    chart.type = "col"
    chart.title = title
    chart.y_axis.title = "Count"
    chart.height, chart.width = 8, 16
    cats = Reference(ws, min_col=col_idx[category_col], min_row=2, max_row=n_rows + 1)
    for vcol in value_cols:
        if vcol not in col_idx:
            continue
        data = Reference(ws, min_col=col_idx[vcol], min_row=1, max_row=n_rows + 1)
        chart.add_data(data, titles_from_data=True)
    if not chart.series:
        return
    chart.set_categories(cats)
    anchor_col = get_column_letter(len(df.columns) + 2)
    ws.add_chart(chart, f"{anchor_col}2")


def _bar_chart(story, title, categories, values, note=None, width=460, height=170,
                bar_color=colors.HexColor("#4A7A9D")):
    """A simple vertical bar chart flowable for the PDF - categories/values
    are same-length parallel lists. Used wherever a report should show a
    breakdown at a glance rather than force the reader to scan a table for
    it (e.g. failures by property)."""
    story.append(Spacer(1, 8))
    story.append(Paragraph(title, STYLES["Heading3"]))
    if note:
        story.append(_p(note))
    if not categories or not any(values):
        story.append(_p("No data recorded."))
        return
    drawing = Drawing(width, height)
    chart = VerticalBarChart()
    chart.x = 45
    chart.y = 30
    chart.height = height - 60
    chart.width = width - 65
    chart.data = [values]
    chart.categoryAxis.categoryNames = [str(c)[:16] for c in categories]
    chart.categoryAxis.labels.angle = 30
    chart.categoryAxis.labels.dy = -12
    chart.categoryAxis.labels.dx = -6
    chart.categoryAxis.labels.fontSize = 7
    chart.valueAxis.valueMin = 0
    chart.bars[0].fillColor = bar_color
    drawing.add(chart)
    story.append(drawing)


def _p(text, style="Normal"):
    """Paragraph with basic XML-escaping, since free-text fields (notes,
    hypotheses, ...) may contain '&', '<', '>' which reportlab's
    Paragraph would otherwise try to interpret as markup."""
    if text is None:
        text = "—"
    escaped = (
        str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )
    return Paragraph(escaped, STYLES[style])


def _key_value_table(pairs, col_widths=(35 * mm, 55 * mm, 35 * mm, 55 * mm)):
    """pairs: list of (label, value) - rendered two-per-row as a compact
    header block (label/value/label/value)."""
    rows = []
    for i in range(0, len(pairs), 2):
        row = []
        for label, value in pairs[i:i + 2]:
            row.extend([label, "—" if value in (None, "") else str(value)])
        while len(row) < 4:
            row.append("")
        rows.append(row)
    t = Table(rows, colWidths=list(col_widths))
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return t


def _section(story, title, rows, col_widths=None):
    """A heading followed by a table built from a list of dicts (all
    sharing the same keys), or a plain "no data" note if rows is empty."""
    story.append(Spacer(1, 8))
    story.append(Paragraph(title, STYLES["Heading3"]))
    if not rows:
        story.append(_p("No data recorded."))
        return
    headers = list(rows[0].keys())
    table_rows = [headers]
    for row in rows:
        table_rows.append(["—" if row.get(h) in (None, "") else str(row.get(h)) for h in headers])
    kwargs = {"colWidths": col_widths} if col_widths else {}
    t = Table(table_rows, repeatRows=1, **kwargs)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6EC")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0BEC5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(t)


def _title_block(story, title, subtitle=None):
    story.append(Paragraph(title, STYLES["Title"]))
    if subtitle:
        story.append(Paragraph(subtitle, STYLES["Normal"]))
    story.append(Spacer(1, 6))


def plant_label(session, plant_id):
    if plant_id is None:
        return "All plants"
    p = session.get(Plant, plant_id)
    return p.name if p else "—"


def product_family_label(session, product_family_id):
    if product_family_id is None:
        return "All product families"
    f = session.get(ProductFamily, product_family_id)
    return f.name if f else "—"


# ---------------------------------------------------------------------------
# 1. Plant / Period Summary Report
# ---------------------------------------------------------------------------

def build_period_summary_data(session, plant_id=None, product_family_id=None, date_from=None, date_to=None, allowed_plant_ids=None):
    """allowed_plant_ids is the tenant-scope guardrail (see tenant_scope.py):
    None = unfiltered (platform owner viewing "All companies"), otherwise the
    list of plant ids the calling company is allowed to see - applied
    unconditionally, on top of whatever single-plant choice plant_id
    represents. Without this, a non-owner user who leaves the on-screen
    "Plant" selector at its default "All plants" would get a report across
    every plant in the database, not just their own company's."""
    runs_q = session.query(ProductionRun)
    if allowed_plant_ids is not None:
        runs_q = runs_q.filter(ProductionRun.plant_id.in_(allowed_plant_ids))
    if plant_id:
        runs_q = runs_q.filter(ProductionRun.plant_id == plant_id)
    if product_family_id:
        runs_q = runs_q.join(FoamGrade, ProductionRun.foam_grade_id == FoamGrade.id).filter(
            FoamGrade.product_family_id == product_family_id
        )
    if date_from:
        runs_q = runs_q.filter(ProductionRun.run_date >= date_from)
    if date_to:
        runs_q = runs_q.filter(ProductionRun.run_date <= date_to)
    runs = runs_q.order_by(ProductionRun.run_date).all()
    run_ids = [r.id for r in runs]

    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.production_run_id.in_(run_ids)).all()
        if run_ids else []
    )
    # Recomputed live rather than trusted from each result's stored
    # pass_fail column - see the same note in
    # analytics.property_results_dataframe.
    computed_verdicts = [compute_pass_fail(r.property_name, r.target_value, r.actual_value) for r in results]
    pass_count = computed_verdicts.count("Pass")
    fail_count = computed_verdicts.count("Fail")
    total_scored = pass_count + fail_count
    pass_rate = round(100 * pass_count / total_scored) if total_scored else None

    observations = (
        session.query(QualityObservation)
        .filter(QualityObservation.production_run_id.in_(run_ids)).all()
        if run_ids else []
    )
    recurring = [o for o in observations if o.frequency == "Recurring"]

    run_rows = [
        {
            "Run ID": r.id, "Date": r.run_date,
            "Foam grade": r.foam_grade.grade_name if r.foam_grade else "—",
            "Recipe version": r.recipe_version.version_label if r.recipe_version else "—",
            "Machine": r.machine.name if r.machine else "—",
            "Batch reference": r.batch_reference or "—",
        }
        for r in runs
    ]

    issue_rows = [
        {
            "Observed": o.observed_at, "Run": o.production_run_id, "Issue type": o.observation_type,
            "Severity": o.severity or "—", "Frequency": o.frequency or "—",
            "Confidence": o.confidence_level or "—",
        }
        for o in observations
    ]

    grade_counts = {}
    for r in runs:
        gname = r.foam_grade.grade_name if r.foam_grade else "—"
        grade_counts[gname] = grade_counts.get(gname, 0) + 1
    grade_breakdown = [{"Foam grade": k, "Production runs": v} for k, v in sorted(grade_counts.items())]

    return {
        "plant": plant_label(session, plant_id),
        "product_family": product_family_label(session, product_family_id),
        "date_from": date_from,
        "date_to": date_to,
        "total_runs": len(runs),
        "pass_rate": pass_rate,
        "total_results_scored": total_scored,
        "total_quality_issues": len(observations),
        "recurring_issues": len(recurring),
        "runs": run_rows,
        "quality_issues": issue_rows,
        "grade_breakdown": grade_breakdown,
    }


def render_period_summary_pdf(data):
    def build(story):
        _title_block(
            story, "Plant / Period Summary Report",
            f"{data['plant']} · {data['product_family']} · {data['date_from'] or 'earliest'} to {data['date_to'] or 'latest'}",
        )
        story.append(_key_value_table([
            ("Production runs", data["total_runs"]),
            ("Quality test pass rate", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—"),
            ("Quality issues", data["total_quality_issues"]),
            ("Recurring quality issues", data["recurring_issues"]),
        ], col_widths=(45 * mm, 40 * mm, 45 * mm, 40 * mm)))
        _section(story, "Production runs in range", data["runs"])
        _section(story, "Quality issues in range", data["quality_issues"])
        _section(story, "Breakdown by foam grade", data["grade_breakdown"])
    return _pdf_bytes(build)


def render_period_summary_excel(data):
    header = [{
        "Plant": data["plant"], "Product family": data["product_family"],
        "Date from": data["date_from"], "Date to": data["date_to"],
        "Production runs": data["total_runs"],
        "Pass rate (%)": data["pass_rate"], "Quality issues": data["total_quality_issues"],
        "Recurring issues": data["recurring_issues"],
    }]
    return _excel_bytes({
        "Summary": header,
        "Production Runs": data["runs"],
        "Quality Issues": data["quality_issues"],
        "Grade Breakdown": data["grade_breakdown"],
    })


# ---------------------------------------------------------------------------
# 2. Trial Closeout Report
#
# Covers the two independent lab-trial flows (see db.CustomerTrial /
# db.OptimizationTrial, added 2026-08-03) - NOT the old TrialRecord concept
# (a formal-experiment flag on a production run), which was removed
# 2026-08-04 after confirming zero real rows across 244 production runs.
# The two models have different closeout fields (a sales-driven Customer
# Trial has customer_name/outcome/customer_feedback; an internally-driven
# Optimization Trial has hypothesis/conclusion/reuse_recommendation), so
# build_trial_report_data() normalizes both into one common "narrative
# fields" list of (label, value) pairs the PDF/Excel renderers can walk
# without needing to know which trial type produced them.
# ---------------------------------------------------------------------------

def build_trial_report_data(session, source_type, trial_id):
    """source_type is "Customer Trial" or "Optimization Trial" (see
    db.SAMPLE_SOURCE_TYPES) - the same source-type string used throughout
    the app to disambiguate the three mutually-exclusive parents a sample/
    quality result can belong to."""
    if source_type == "Customer Trial":
        trial = session.get(CustomerTrial, trial_id)
    elif source_type == "Optimization Trial":
        trial = session.get(OptimizationTrial, trial_id)
    else:
        return None
    if trial is None:
        return None
    grade = trial.foam_grade
    plant = trial.plant

    quality_issues = [
        {
            "Issue type": o.observation_type, "Severity": o.severity or "—",
            "Frequency": o.frequency or "—", "Confidence": o.confidence_level or "—",
        }
        for o in session.query(QualityObservation).filter(
            (QualityObservation.customer_trial_id == trial_id)
            if source_type == "Customer Trial"
            else (QualityObservation.optimization_trial_id == trial_id)
        ).all()
    ]

    if source_type == "Customer Trial":
        narrative_fields = [
            ("Customer", trial.customer_name),
            ("Sales opportunity reference", trial.sales_opportunity_reference or "—"),
            ("Requested by", trial.requested_by or "—"),
            ("Trial objective", trial.trial_objective or "—"),
            ("Outcome", trial.outcome or "—"),
            ("Customer feedback", trial.customer_feedback or "—"),
            ("Follow-up action", trial.follow_up_action or "—"),
        ]
    else:
        narrative_fields = [
            ("Improvement initiative reference", trial.improvement_initiative_reference or "—"),
            ("Hypothesis", trial.hypothesis or "—"),
            ("What changed", trial.what_changed or "—"),
            ("Result against target", trial.result_against_target or "—"),
            ("Conclusion", trial.conclusion or "—"),
            ("Reuse recommendation", trial.reuse_recommendation or "—"),
        ]

    return {
        "source_type": source_type,
        "trial_id": trial.id,
        "foam_grade": grade.grade_name if grade else "—",
        "plant": plant.name if plant else "—",
        "status": trial.status,
        "responsible_person": trial.responsible_person or "—",
        "trial_date": trial.trial_date,
        "batch_reference": trial.batch_reference or "—",
        "notes": trial.notes or "",
        "narrative_fields": narrative_fields,
        "reviewed_by": trial.reviewed_by or "—",
        # Only OptimizationTrial has a separate approved_by (CustomerTrial's
        # closeout is reviewed_by only - see db.py's REQUIRED_CLOSEOUT_FIELDS
        # on each model).
        "approved_by": getattr(trial, "approved_by", None) or "—",
        "date_closed": trial.date_closed,
        "quality_issues": quality_issues,
    }


def render_trial_report_pdf(data):
    def build(story):
        _title_block(
            story, f"Trial Closeout Report — {data['source_type']} #{data['trial_id']}",
            f"{data['foam_grade']} · {data['plant']} · {data['status']}",
        )
        story.append(_key_value_table([
            ("Status", data["status"]), ("Responsible", data["responsible_person"]),
            ("Foam grade", data["foam_grade"]), ("Plant", data["plant"]),
            ("Trial date", data["trial_date"]), ("Batch reference", data["batch_reference"]),
            ("Reviewed by", data["reviewed_by"]), ("Approved by", data["approved_by"]),
            ("Date closed", data["date_closed"]), ("", ""),
        ]))
        if data["notes"]:
            story.append(Spacer(1, 6))
            story.append(_p(f"Notes: {data['notes']}"))
        story.append(Spacer(1, 8))
        for label, value in data["narrative_fields"]:
            story.append(Paragraph(label, STYLES["Heading3"]))
            story.append(_p(value))
            story.append(Spacer(1, 6))
        _section(story, "Quality issues observed", data["quality_issues"])
    return _pdf_bytes(build)


def render_trial_report_excel(data):
    header = [{
        "Trial type": data["source_type"], "Trial ID": data["trial_id"],
        "Foam grade": data["foam_grade"], "Plant": data["plant"],
        "Status": data["status"], "Responsible": data["responsible_person"],
        "Trial date": data["trial_date"], "Batch reference": data["batch_reference"],
        **{label: value for label, value in data["narrative_fields"]},
        "Reviewed by": data["reviewed_by"], "Approved by": data["approved_by"],
        "Date closed": data["date_closed"], "Notes": data["notes"],
    }]
    return _excel_bytes({
        "Trial": header,
        "Quality Issues": data["quality_issues"],
    })


# ---------------------------------------------------------------------------
# 3. Recipe / Formulation Record Report
#
# Internal-use record for one recipe version: the formulation, its quality
# specs vs. actual results aggregated over a chosen date range, and cost
# per kg. NOT customer-facing as-is (see reports.py module docstring) -
# the recipe/formulation section is the whole reason it can't be handed
# to a customer.
#
# "Quality specs & results" pulls actual PhysicalPropertyResult rows from
# every production run, customer trial, and optimization trial built on
# this recipe version (the three mutually-exclusive parents - see
# db.SAMPLE_SOURCE_TYPES), filtered to the caller's date range, then
# aggregated per property (average actual, pass rate, sample count)
# against that foam grade's target - never a flat row-by-row dump, per
# the Reports redesign ruling that raw data belongs in each page's own
# CSV export, not in a report.
# ---------------------------------------------------------------------------

def _recipe_version_target_properties(grade):
    """FoamGrade's target specs as a flat list of {property_name, target_value,
    unit} dicts - density/hardness (fixed columns on FoamGrade) plus any
    additional targets recorded in FoamGradeTargetProperty."""
    if grade is None:
        return []
    targets = []
    if grade.target_density is not None:
        targets.append({"property_name": "Density", "target_value": grade.target_density, "unit": "kg/m³"})
    if grade.target_hardness is not None:
        # "40% IFD / hardness" is the canonical property_name used app-wide
        # (see quality_standards.INDUSTRY_TOLERANCES and pages/15_Recipe_
        # Optimization.py's own target_by_name dict) - matching it here,
        # rather than a differently-worded label, is what lets this target
        # line up with actual PhysicalPropertyResult rows recorded against
        # that same property name instead of showing as two separate rows.
        targets.append({"property_name": "40% IFD / hardness", "target_value": grade.target_hardness, "unit": "N"})
    for tp in grade.target_properties:
        targets.append({"property_name": tp.property_name, "target_value": tp.target_value, "unit": tp.unit or ""})
    return targets


def _property_results_for_recipe_version(session, recipe_version_id, date_from=None, date_to=None):
    """Every PhysicalPropertyResult tied (via production run, customer trial,
    or optimization trial - see db.SAMPLE_SOURCE_TYPES) to this recipe
    version, filtered to [date_from, date_to] on that parent's own date
    field (run_date / trial_date). Three separate joins rather than one
    query, since which date field applies depends on which of the three
    parent types the result belongs to."""
    results = []

    run_q = (
        session.query(PhysicalPropertyResult)
        .join(ProductionRun, PhysicalPropertyResult.production_run_id == ProductionRun.id)
        .filter(ProductionRun.recipe_version_id == recipe_version_id)
    )
    if date_from:
        run_q = run_q.filter(ProductionRun.run_date >= date_from)
    if date_to:
        run_q = run_q.filter(ProductionRun.run_date <= date_to)
    results.extend(run_q.all())

    ct_q = (
        session.query(PhysicalPropertyResult)
        .join(CustomerTrial, PhysicalPropertyResult.customer_trial_id == CustomerTrial.id)
        .filter(CustomerTrial.recipe_version_id == recipe_version_id)
    )
    if date_from:
        ct_q = ct_q.filter(CustomerTrial.trial_date >= date_from)
    if date_to:
        ct_q = ct_q.filter(CustomerTrial.trial_date <= date_to)
    results.extend(ct_q.all())

    ot_q = (
        session.query(PhysicalPropertyResult)
        .join(OptimizationTrial, PhysicalPropertyResult.optimization_trial_id == OptimizationTrial.id)
        .filter(OptimizationTrial.recipe_version_id == recipe_version_id)
    )
    if date_from:
        ot_q = ot_q.filter(OptimizationTrial.trial_date >= date_from)
    if date_to:
        ot_q = ot_q.filter(OptimizationTrial.trial_date <= date_to)
    results.extend(ot_q.all())

    return results


def build_recipe_formulation_record_data(session, recipe_version_id, date_from=None, date_to=None):
    rv = session.get(RecipeVersion, recipe_version_id)
    if rv is None:
        return None
    grade = rv.foam_grade
    family = grade.product_family if grade else None

    ordered_components = sorted(
        rv.components,
        key=lambda c: (c.role_in_formulation or "", c.raw_material_name or ""),
    )
    components = [
        {
            "Material": c.raw_material_name,
            "Supplier": c.supplier or "—",
            "PHP": c.php,
            "Role": c.role_in_formulation or "—",
            "Notes": c.notes or "—",
        }
        for c in ordered_components
    ]

    targets_by_name = {t["property_name"]: t for t in _recipe_version_target_properties(grade)}
    results = _property_results_for_recipe_version(session, rv.id, date_from, date_to)
    by_property = {}
    for r in results:
        by_property.setdefault(r.property_name, []).append(r)

    quality_rows = []
    # Show every declared target even if zero results fell in range (an
    # honest "no data yet" beats silently omitting the row), plus any
    # measured property that isn't a formally declared target.
    for prop_name in sorted(set(targets_by_name) | set(by_property)):
        rs = by_property.get(prop_name, [])
        verdicts = [compute_pass_fail(r.property_name, r.target_value, r.actual_value) for r in rs]
        pass_ct, fail_ct = verdicts.count("Pass"), verdicts.count("Fail")
        scored = pass_ct + fail_ct
        actuals = [r.actual_value for r in rs if r.actual_value is not None]
        target = targets_by_name.get(prop_name)
        target_value = target["target_value"] if target else next(
            (r.target_value for r in rs if r.target_value is not None), None
        )
        unit = (target["unit"] if target else None) or next((r.unit for r in rs if r.unit), "")
        quality_rows.append({
            "Property": prop_name,
            "Target": target_value,
            "Avg. actual": round(sum(actuals) / len(actuals), 2) if actuals else None,
            "Unit": unit,
            "Results in range": len(rs),
            "Pass rate": f"{round(100 * pass_ct / scored)}%" if scored else "—",
        })

    cost = recipe_version_cost(session, rv)
    # Same php-parts-as-kg convention as pages/15_Recipe_Optimization.py's
    # _cost_per_kg() - a formulation's php total already IS its cost basis
    # (see analytics.recipe_version_cost's own docstring), so cost per kg
    # is simply total cost / total php, once any component is priced.
    cost_per_kg = (
        round(cost["total_cost"] / cost["total_php"], 2)
        if cost["total_cost"] is not None and cost["total_php"]
        else None
    )

    return {
        "recipe_version_id": rv.id,
        "version_label": rv.version_label,
        "foam_grade": grade.grade_name if grade else "—",
        "product_family": family.name if family else "—",
        "approval_status": rv.approval_status,
        "is_active": rv.is_active,
        "effective_date": rv.effective_date,
        "created_by": rv.created_by or "—",
        "change_note": rv.change_note or "",
        "ratio_index": rv.ratio_index,
        "components": components,
        "date_from": date_from,
        "date_to": date_to,
        "quality_rows": quality_rows,
        "cost_per_kg": cost_per_kg,
        "cost_priced_php": cost["priced_php"],
        "cost_total_php": cost["total_php"],
        "cost_missing_materials": cost["missing"],
    }


def render_recipe_formulation_record_pdf(data):
    def build(story):
        _title_block(
            story, f"Recipe / Formulation Record — {data['version_label']}",
            f"{data['foam_grade']} · {data['product_family']} · "
            f"{'Active recipe' if data['is_active'] else 'Retired version'}",
        )
        story.append(_key_value_table([
            ("Approval status", data["approval_status"]), ("Active", "Yes" if data["is_active"] else "No"),
            ("Effective date", data["effective_date"]), ("Created by", data["created_by"]),
            ("Ratio / index", f"{data['ratio_index']:.3f}" if data["ratio_index"] is not None else "—"),
            ("", ""),
        ]))
        if data["change_note"]:
            story.append(Spacer(1, 6))
            story.append(_p(f"Change note: {data['change_note']}"))
        _section(story, "Formulation (recipe components)", data["components"])

        date_range = f"{data['date_from'] or 'earliest'} to {data['date_to'] or 'latest'}"
        _section(story, f"Quality specs vs. results ({date_range})", data["quality_rows"])

        story.append(Spacer(1, 8))
        story.append(Paragraph("Cost", STYLES["Heading3"]))
        if data["cost_per_kg"] is None:
            story.append(_p("No priced components - cost per kg cannot be calculated."))
        else:
            coverage = f"{data['cost_priced_php']:.2f} / {data['cost_total_php']:.2f} php priced"
            story.append(_key_value_table([
                ("Cost per kg", data["cost_per_kg"]), ("Cost coverage", coverage),
            ]))
            if data["cost_missing_materials"]:
                story.append(_p("Unpriced materials (excluded from total): " + ", ".join(data["cost_missing_materials"])))
    return _pdf_bytes(build)


def render_recipe_formulation_record_excel(data):
    header = [{
        "Recipe version": data["version_label"], "Foam grade": data["foam_grade"],
        "Product family": data["product_family"], "Approval status": data["approval_status"],
        "Active": "Yes" if data["is_active"] else "No", "Effective date": data["effective_date"],
        "Created by": data["created_by"], "Ratio / index": data["ratio_index"],
        "Change note": data["change_note"], "Date from": data["date_from"], "Date to": data["date_to"],
        "Cost per kg": data["cost_per_kg"],
        "Cost coverage (php priced / total)": f"{data['cost_priced_php']} / {data['cost_total_php']}",
        "Unpriced materials": ", ".join(data["cost_missing_materials"]) or "—",
    }]
    return _excel_bytes({
        "Header": header,
        "Formulation": data["components"],
        "Quality vs Spec": data["quality_rows"],
    })


# ---------------------------------------------------------------------------
# 4. Where Used Report
#
# Given a raw material, answers "which recipes use this, and what depends
# on it" - the reverse lookup a Plant Manager needs before considering a
# material substitution. Scoped inherently by the tenant boundary: it
# joins on RecipeComponent.raw_material_id, a real FK to one specific
# (already company-scoped) raw_materials row, not a name match, so it
# can't cross into another company's recipes.
# ---------------------------------------------------------------------------

def build_where_used_report_data(session, raw_material_id):
    rm = session.get(RawMaterial, raw_material_id)
    if rm is None:
        return None

    components = (
        session.query(RecipeComponent)
        .filter(RecipeComponent.raw_material_id == rm.id)
        .all()
    )
    recipe_version_ids = {c.recipe_version_id for c in components}
    versions = (
        session.query(RecipeVersion).filter(RecipeVersion.id.in_(recipe_version_ids)).all()
        if recipe_version_ids else []
    )
    version_by_id = {v.id: v for v in versions}

    def _sort_key(c):
        v = version_by_id.get(c.recipe_version_id)
        grade = v.foam_grade if v else None
        return (grade.grade_name if grade else "", v.version_label if v else "")

    usage_rows = []
    grade_ids, family_names = set(), set()
    for c in sorted(components, key=_sort_key):
        v = version_by_id.get(c.recipe_version_id)
        grade = v.foam_grade if v else None
        family = grade.product_family if grade else None
        if grade:
            grade_ids.add(grade.id)
        if family:
            family_names.add(family.name)
        usage_rows.append({
            "Foam grade": grade.grade_name if grade else "—",
            "Product family": family.name if family else "—",
            "Recipe version": v.version_label if v else "—",
            "Status": "Active" if v and v.is_active else "Retired",
            "PHP": c.php,
            "Role": c.role_in_formulation or "—",
            "Approval status": v.approval_status if v else "—",
        })

    grades = session.query(FoamGrade).filter(FoamGrade.id.in_(grade_ids)).all() if grade_ids else []
    target_rows = []
    for g in sorted(grades, key=lambda g: g.grade_name):
        for t in _recipe_version_target_properties(g):
            target_rows.append({
                "Foam grade": g.grade_name, "Property": t["property_name"],
                "Target": t["target_value"], "Unit": t["unit"],
            })

    trial_rows = []
    if recipe_version_ids:
        customer_trials = (
            session.query(CustomerTrial)
            .filter(CustomerTrial.recipe_version_id.in_(recipe_version_ids))
            .order_by(CustomerTrial.trial_date.desc())
            .all()
        )
        for t in customer_trials:
            trial_rows.append({
                "Trial type": "Customer Trial", "Trial ID": t.id,
                "Foam grade": t.foam_grade.grade_name if t.foam_grade else "—",
                "Status": t.status, "Trial date": t.trial_date, "Outcome": t.outcome or "—",
            })
        optimization_trials = (
            session.query(OptimizationTrial)
            .filter(OptimizationTrial.recipe_version_id.in_(recipe_version_ids))
            .order_by(OptimizationTrial.trial_date.desc())
            .all()
        )
        for t in optimization_trials:
            trial_rows.append({
                "Trial type": "Optimization Trial", "Trial ID": t.id,
                "Foam grade": t.foam_grade.grade_name if t.foam_grade else "—",
                "Status": t.status, "Trial date": t.trial_date, "Outcome": t.conclusion or "—",
            })

    return {
        "raw_material_id": rm.id,
        "raw_material_name": rm.name,
        "category": rm.category or "—",
        "default_supplier": rm.default_supplier or "—",
        "active": rm.active,
        "recipe_version_count": len(recipe_version_ids),
        "foam_grade_count": len(grade_ids),
        "product_family_count": len(family_names),
        "usage_rows": usage_rows,
        "target_rows": target_rows,
        "trial_rows": trial_rows,
    }


def render_where_used_report_pdf(data):
    def build(story):
        _title_block(
            story, f"Where Used Report — {data['raw_material_name']}",
            f"{data['category']} · Default supplier: {data['default_supplier']} · "
            f"{'Active' if data['active'] else 'Inactive'} material",
        )
        story.append(_key_value_table([
            ("Recipe versions using this material", data["recipe_version_count"]),
            ("Foam grades affected", data["foam_grade_count"]),
            ("Product families affected", data["product_family_count"]),
            ("", ""),
        ]))
        _section(story, "Recipes using this material", data["usage_rows"])
        _section(story, "Target properties of affected foam grades", data["target_rows"])
        _section(story, "Trial precedent (Customer / Optimization Trials on these recipes)", data["trial_rows"])
    return _pdf_bytes(build)


def render_where_used_report_excel(data):
    header = [{
        "Raw material": data["raw_material_name"], "Category": data["category"],
        "Default supplier": data["default_supplier"], "Active": "Yes" if data["active"] else "No",
        "Recipe versions using it": data["recipe_version_count"],
        "Foam grades affected": data["foam_grade_count"],
        "Product families affected": data["product_family_count"],
    }]
    return _excel_bytes({
        "Header": header,
        "Recipe Usage": data["usage_rows"],
        "Target Properties": data["target_rows"],
        "Trial Precedent": data["trial_rows"],
    })


# ---------------------------------------------------------------------------
# 5. Batch Release / Conformance Record (Production Run)
#
# Purpose (per user direction 2026-08-04): "did this batch meet spec, and
# is there anything on record I should know" - not a transcription of
# every stored field on the run. A clean run gets a short document,
# header + recipe + verdict. A FLAGGED run (a failed quality result, or
# any recorded quality issue) widens to pull relevant context from every
# other tab on the Production Run page - Setup-vs-Finalized process
# settings, actual component stream readings, and Production Events -
# because a flag's explanation often isn't on the same tab as the flag
# itself.
# ---------------------------------------------------------------------------

# Tolerance for treating two ProductionPhase float settings as "the same"
# rather than flagging a deviation - avoids surfacing meaningless float
# noise (25.000001 vs 25.0) as if it were a real setpoint change.
_SETTING_DEVIATION_EPSILON = 0.01
_FALLPLATE_POSITION_DEVIATION_MM = 2.0  # per db.py: fall-plate position materially affects density profile


def _phase_by_name(phases, name):
    return next((p for p in phases if p.phase_name == name), None)


def _setup_vs_finalized_deviations(session, setup_phase, finalized_phase):
    """Every process setting that actually changed between the Setup
    (planned) and Finalized (actual) phase snapshots of one run - not the
    full settings table (see PHASE_SETTING_FIELDS in analytics.py), just
    the rows that differ, since those are what could explain a flagged
    result. Includes foaming_mode (a controlled-vocabulary field, not part
    of PHASE_SETTING_FIELDS) alongside the numeric settings."""
    if setup_phase is None or finalized_phase is None:
        return []
    deviations = []
    for field in PHASE_SETTING_FIELDS:
        setup_val = getattr(setup_phase, field)
        final_val = getattr(finalized_phase, field)
        if setup_val is None and final_val is None:
            continue
        if (
            setup_val is not None and final_val is not None
            and abs(float(setup_val) - float(final_val)) <= _SETTING_DEVIATION_EPSILON
        ):
            continue
        if setup_val == final_val:
            continue
        deviations.append({
            "Setting": PHASE_SETTING_LABELS.get(field, field),
            "Setup (planned)": setup_val, "Finalized (actual)": final_val,
        })
    if setup_phase.foaming_mode != finalized_phase.foaming_mode:
        deviations.append({
            "Setting": "Foaming mode",
            "Setup (planned)": setup_phase.foaming_mode or "—",
            "Finalized (actual)": finalized_phase.foaming_mode or "—",
        })
    return deviations


def _fallplate_deviations(session, setup_phase, finalized_phase):
    """Fall-plate section position changes between Setup and Finalized,
    keyed by section_number - only rows whose position moved more than
    _FALLPLATE_POSITION_DEVIATION_MM, the same "only what changed"
    principle as _setup_vs_finalized_deviations."""
    if setup_phase is None or finalized_phase is None:
        return []
    setup_by_section = {
        p.section_number: p
        for p in session.query(FallplateSectionPosition)
        .filter(FallplateSectionPosition.production_phase_id == setup_phase.id).all()
    }
    final_by_section = {
        p.section_number: p
        for p in session.query(FallplateSectionPosition)
        .filter(FallplateSectionPosition.production_phase_id == finalized_phase.id).all()
    }
    deviations = []
    for section in sorted(set(setup_by_section) | set(final_by_section)):
        s, f = setup_by_section.get(section), final_by_section.get(section)
        s_pos = s.position_mm if s else None
        f_pos = f.position_mm if f else None
        if s_pos is None or f_pos is None or abs(s_pos - f_pos) <= _FALLPLATE_POSITION_DEVIATION_MM:
            continue
        deviations.append({
            "Section": section, "Setup position (mm)": s_pos, "Finalized position (mm)": f_pos,
            "Change (mm)": round(f_pos - s_pos, 1),
        })
    return deviations


def build_batch_release_record_data(session, run_id):
    run = session.get(ProductionRun, run_id)
    if run is None:
        return None
    grade = run.foam_grade
    family = grade.product_family if grade else None
    recipe = run.recipe_version

    ordered_components = (
        sorted(recipe.components, key=lambda c: (c.role_in_formulation or "", c.raw_material_name or ""))
        if recipe else []
    )
    recipe_components = [
        {
            "Material": c.raw_material_name, "Supplier": c.supplier or "—",
            "PHP": c.php, "Role": c.role_in_formulation or "—", "Notes": c.notes or "—",
        }
        for c in ordered_components
    ]

    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.production_run_id == run_id).all()
    )
    quality_results = [
        {
            "Property": r.property_name, "Target": r.target_value, "Actual": r.actual_value,
            "Unit": r.unit or "",
            "Pass/Fail": compute_pass_fail(r.property_name, r.target_value, r.actual_value) or "—",
            "Tested": r.tested_at,
        }
        for r in results
    ]
    verdicts = [compute_pass_fail(r.property_name, r.target_value, r.actual_value) for r in results]
    if not results:
        quality_verdict = "No testing recorded"
    elif "Fail" in verdicts:
        quality_verdict = "Non-conforming"
    elif verdicts and all(v == "Pass" for v in verdicts):
        quality_verdict = "Conforming"
    else:
        quality_verdict = "Incomplete testing"

    quality_issues = [
        {
            "Issue type": o.observation_type, "Severity": o.severity or "—",
            "Frequency": o.frequency or "—", "Confidence": o.confidence_level or "—",
            "Suspected cause": o.suspected_cause or "—",
        }
        for o in session.query(QualityObservation).filter(QualityObservation.production_run_id == run_id).all()
    ]

    has_flags = ("Fail" in verdicts) or bool(quality_issues)
    flag_reasons = []
    fail_count = verdicts.count("Fail")
    if fail_count:
        flag_reasons.append(f"{fail_count} failed quality result(s)")
    if quality_issues:
        flag_reasons.append(f"{len(quality_issues)} quality issue(s) recorded")

    setup_deviations, stream_readings, stream_calibration_flags, production_events, fallplate_deviations = (
        [], [], [], [], [],
    )
    if has_flags:
        phases = session.query(ProductionPhase).filter(ProductionPhase.production_run_id == run_id).all()
        setup_phase = _phase_by_name(phases, "Setup")
        finalized_phase = _phase_by_name(phases, "Finalized")
        setup_deviations = _setup_vs_finalized_deviations(session, setup_phase, finalized_phase)
        fallplate_deviations = _fallplate_deviations(session, setup_phase, finalized_phase)

        if finalized_phase is not None:
            readings = (
                session.query(ComponentStreamReading)
                .filter(ComponentStreamReading.production_phase_id == finalized_phase.id).all()
            )
            stream_readings = [
                {
                    "Stream": rd.stream_name, "Flow": rd.flow, "Unit": rd.flow_unit or "",
                    "Pump speed": rd.pump_speed, "Total delivered": rd.flow_total_qty,
                    "Temperature (°C)": rd.temperature_c, "Pressure (bar)": rd.pressure_bar,
                    "Calibration": rd.calibration_status or "—",
                }
                for rd in readings
            ]
            stream_calibration_flags = [
                rd.stream_name for rd in readings if rd.calibration_status and rd.calibration_status != "Valid"
            ]

        production_events = [
            {
                "Time": e.event_ts, "Type": e.event_type, "Severity": e.severity or "—",
                "Description": e.description or "—",
            }
            for e in session.query(ProductionEvent)
            .filter(ProductionEvent.production_run_id == run_id)
            .order_by(ProductionEvent.event_ts).all()
        ]

    return {
        "run_id": run.id,
        "plant": run.plant.name if run.plant else "—",
        "product_family": family.name if family else "—",
        "foam_grade": grade.grade_name if grade else "—",
        "machine": run.machine.name if run.machine else "—",
        "run_date": run.run_date,
        "batch_reference": run.batch_reference or "—",
        "block_reference": run.block_reference or "—",
        "operator": run.operator_or_team_reference or "—",
        "notes": run.notes or "",
        "recipe_version_label": recipe.version_label if recipe else "—",
        "recipe_approval_status": recipe.approval_status if recipe else "—",
        "recipe_effective_date": recipe.effective_date if recipe else None,
        "recipe_ratio_index": recipe.ratio_index if recipe else None,
        "recipe_components": recipe_components,
        "quality_results": quality_results,
        "quality_verdict": quality_verdict,
        "quality_issues": quality_issues,
        "has_flags": has_flags,
        "flag_reasons": flag_reasons,
        "setup_deviations": setup_deviations,
        "fallplate_deviations": fallplate_deviations,
        "stream_readings": stream_readings,
        "stream_calibration_flags": stream_calibration_flags,
        "production_events": production_events,
    }


def render_batch_release_record_pdf(data):
    def build(story):
        _title_block(
            story, f"Batch Release Record — Run #{data['run_id']}",
            f"{data['plant']} · {data['foam_grade']} · {data['run_date'] or '—'} · "
            f"Verdict: {data['quality_verdict']}",
        )
        story.append(_key_value_table([
            ("Plant", data["plant"]), ("Product family", data["product_family"]),
            ("Foam grade", data["foam_grade"]), ("Machine", data["machine"]),
            ("Run date", data["run_date"]), ("Batch reference", data["batch_reference"]),
            ("Block reference", data["block_reference"]), ("Operator/team", data["operator"]),
            ("Quality verdict", data["quality_verdict"]), ("", ""),
        ]))
        if data["notes"]:
            story.append(Spacer(1, 6))
            story.append(_p(f"Notes: {data['notes']}"))

        story.append(Spacer(1, 8))
        story.append(Paragraph("Recipe used", STYLES["Heading3"]))
        story.append(_key_value_table([
            ("Recipe version", data["recipe_version_label"]), ("Approval status", data["recipe_approval_status"]),
            ("Effective date", data["recipe_effective_date"]),
            ("Ratio / index", f"{data['recipe_ratio_index']:.3f}" if data["recipe_ratio_index"] is not None else "—"),
        ]))
        _section(story, "Formulation", data["recipe_components"])

        _section(story, "Quality test results", data["quality_results"])
        _section(story, "Quality issues", data["quality_issues"])

        if data["has_flags"]:
            story.append(Spacer(1, 10))
            story.append(Paragraph("Flagged — supporting context from other tabs", STYLES["Heading2"]))
            story.append(_p("Flagged because: " + "; ".join(data["flag_reasons"])))
            _section(story, "Process setting changes (Setup → Finalized)", data["setup_deviations"])
            _section(story, "Fall-plate position changes (Setup → Finalized)", data["fallplate_deviations"])
            _section(story, "Component stream readings (Finalized phase)", data["stream_readings"])
            if data["stream_calibration_flags"]:
                story.append(_p(
                    "⚠ Non-valid calibration status recorded for: " + ", ".join(data["stream_calibration_flags"])
                ))
            _section(story, "Production events during this run", data["production_events"])
    return _pdf_bytes(build)


def render_batch_release_record_excel(data):
    header = [{
        "Run ID": data["run_id"], "Plant": data["plant"], "Product family": data["product_family"],
        "Foam grade": data["foam_grade"], "Machine": data["machine"], "Run date": data["run_date"],
        "Batch reference": data["batch_reference"], "Block reference": data["block_reference"],
        "Operator/team": data["operator"], "Quality verdict": data["quality_verdict"],
        "Recipe version": data["recipe_version_label"], "Recipe approval status": data["recipe_approval_status"],
        "Notes": data["notes"],
        "Flagged": "Yes" if data["has_flags"] else "No",
        "Flag reasons": "; ".join(data["flag_reasons"]) if data["flag_reasons"] else "—",
    }]
    sheets = {
        "Header": header,
        "Formulation": data["recipe_components"],
        "Quality Results": data["quality_results"],
        "Quality Issues": data["quality_issues"],
    }
    if data["has_flags"]:
        sheets["Process Setting Changes"] = data["setup_deviations"]
        sheets["Fallplate Changes"] = data["fallplate_deviations"]
        sheets["Stream Readings"] = data["stream_readings"]
        sheets["Production Events"] = data["production_events"]
    return _excel_bytes(sheets)


# ---------------------------------------------------------------------------
# 6. Quality Test Result Report (Physical Property Result page)
#
# Placement (per user direction 2026-08-04): this report's subject is a
# comprehensive multi-field selection (Pass/Fail, Property, Foam scope) the
# reader has to build up first, not a single dropdown choice - so it lives
# on pages/5_Physical_Property_Result.py itself, right below the same
# filter controls and Pareto chart it shares its scope with, rather than
# on the Report page. build_quality_test_report_data() never re-derives
# tenant scope or filters on its own - it purely aggregates the exact
# PhysicalPropertyResult id set the page has already scoped and filtered,
# so the report always matches what's on screen at the moment it's
# generated.
# ---------------------------------------------------------------------------

def _qtr_source_and_grade(result):
    """(source label, human-readable parent description, foam grade name)
    for a PhysicalPropertyResult, resolving whichever of the three
    mutually exclusive parents (production run / customer trial /
    optimization trial - see db.SAMPLE_SOURCE_TYPES) it belongs to.
    Mirrors pages/5_Physical_Property_Result.py's own
    _result_source_desc()/_result_foam_grade_id() - kept as a local copy
    here since reports.py doesn't import from page modules."""
    if result.production_run_id is not None:
        run = result.production_run
        grade = run.foam_grade if run else None
        desc = f"Run #{run.id} — {grade.grade_name if grade else '—'} · {run.run_date}" if run else f"Run #{result.production_run_id}"
        return "Production Run", desc, grade.grade_name if grade else "—"
    if result.customer_trial_id is not None:
        t = result.customer_trial
        grade = t.foam_grade if t else None
        desc = f"Trial #{t.id} — {t.customer_name}" if t else f"Trial #{result.customer_trial_id}"
        return "Customer Trial", desc, grade.grade_name if grade else "—"
    if result.optimization_trial_id is not None:
        t = result.optimization_trial
        grade = t.foam_grade if t else None
        ref = (t.improvement_initiative_reference or "(no reference)") if t else ""
        desc = f"Trial #{t.id} — {ref}" if t else f"Trial #{result.optimization_trial_id}"
        return "Optimization Trial", desc, grade.grade_name if grade else "—"
    return "—", "—", "—"


def build_quality_test_report_data(session, result_ids, scope):
    """result_ids: PhysicalPropertyResult ids already scoped (tenant) and
    filtered (Pass/Fail, Property, Foam scope) by the caller - see the
    module-level note above. scope: dict of already-formatted display
    strings describing what was selected - pass_fail_label,
    property_label, foam_scope_label - shown in the report header so the
    reader knows exactly what this report does and doesn't cover."""
    results = (
        session.query(PhysicalPropertyResult)
        .filter(PhysicalPropertyResult.id.in_(result_ids)).all()
        if result_ids else []
    )

    detail = []
    for r in results:
        verdict = compute_pass_fail(r.property_name, r.target_value, r.actual_value) or "Not computed"
        source_label, source_desc, grade_name = _qtr_source_and_grade(r)
        detail.append({
            "result": r, "verdict": verdict,
            "source_desc": source_desc, "grade_name": grade_name,
        })

    total = len(detail)
    pass_count = sum(1 for d in detail if d["verdict"] == "Pass")
    fail_count = sum(1 for d in detail if d["verdict"] == "Fail")
    not_computed_count = total - pass_count - fail_count
    total_scored = pass_count + fail_count
    pass_rate = round(100 * pass_count / total_scored) if total_scored else None

    # Failures by property - the same grouping as the on-page Pareto chart
    # (helpers.render_pareto_chart), recomputed here since reports.py
    # can't import a Streamlit-rendering helper. Fail-only, not
    # pass+fail+not-computed per property, since "which properties are
    # behind the failures" is the actionable question a Pareto answers.
    property_fail_counts = {}
    for d in detail:
        if d["verdict"] == "Fail":
            prop = d["result"].property_name
            property_fail_counts[prop] = property_fail_counts.get(prop, 0) + 1
    property_breakdown = [
        {"Property": k, "Fail count": v}
        for k, v in sorted(property_fail_counts.items(), key=lambda kv: -kv[1])
    ]

    grade_counts = {}
    for d in detail:
        bucket = grade_counts.setdefault(d["grade_name"], {"Pass": 0, "Fail": 0, "Not computed": 0})
        bucket[d["verdict"]] += 1
    grade_breakdown = [
        {"Foam grade": g, "Pass count": c["Pass"], "Fail count": c["Fail"]}
        for g, c in sorted(grade_counts.items())
    ]
    # Only meaningful as a chart when the selection actually spans more
    # than one grade - a single-grade selection would just re-draw the
    # header metrics as a one-bar "chart".
    show_grade_breakdown = len(grade_counts) > 1

    failing_results = [
        {
            "Source": d["source_desc"], "Property": d["result"].property_name,
            "Target": d["result"].target_value, "Actual": d["result"].actual_value,
            "Unit": d["result"].unit or "",
            "Deviation": (
                round(d["result"].actual_value - d["result"].target_value, 2)
                if d["result"].actual_value is not None and d["result"].target_value is not None
                else None
            ),
            "Foam grade": d["grade_name"], "Tested": d["result"].tested_at,
        }
        for d in sorted(detail, key=lambda d: (d["result"].property_name, d["source_desc"]))
        if d["verdict"] == "Fail"
    ]

    return {
        "scope": scope,
        "total_results": total,
        "pass_count": pass_count,
        "fail_count": fail_count,
        "not_computed_count": not_computed_count,
        "pass_rate": pass_rate,
        "property_breakdown": property_breakdown,
        "grade_breakdown": grade_breakdown,
        "show_grade_breakdown": show_grade_breakdown,
        "failing_results": failing_results,
    }


def render_quality_test_report_pdf(data):
    scope = data["scope"]

    def build(story):
        _title_block(
            story, "Quality Test Result Report",
            f"Pass/Fail: {scope['pass_fail_label']} · Property: {scope['property_label']} · "
            f"Foam scope: {scope['foam_scope_label']}",
        )
        story.append(_key_value_table([
            ("Results", data["total_results"]),
            ("Pass rate", f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—"),
            ("Pass", data["pass_count"]), ("Fail", data["fail_count"]),
            ("Not computed", data["not_computed_count"]), ("", ""),
        ]))

        _bar_chart(
            story, "Pass / Fail breakdown",
            ["Pass", "Fail", "Not computed"],
            [data["pass_count"], data["fail_count"], data["not_computed_count"]],
        )

        prop_rows = data["property_breakdown"]
        _bar_chart(
            story, "Failures by property",
            [row["Property"] for row in prop_rows], [row["Fail count"] for row in prop_rows],
            note="Which tested properties are behind the failures in this selection." if prop_rows else None,
        )

        if data["show_grade_breakdown"]:
            grade_rows = data["grade_breakdown"]
            _bar_chart(
                story, "Failures by foam grade",
                [row["Foam grade"] for row in grade_rows], [row["Fail count"] for row in grade_rows],
                note="Shown because this selection spans more than one foam grade.",
            )

        _section(story, "Failing results (target vs. actual)", data["failing_results"])
    return _pdf_bytes(build)


def render_quality_test_report_excel(data):
    scope = data["scope"]
    header = [{
        "Pass/Fail filter": scope["pass_fail_label"], "Property filter": scope["property_label"],
        "Foam scope": scope["foam_scope_label"], "Results in selection": data["total_results"],
        "Pass rate": f"{data['pass_rate']}%" if data["pass_rate"] is not None else "—",
        "Pass count": data["pass_count"], "Fail count": data["fail_count"],
        "Not computed count": data["not_computed_count"],
    }]
    pass_fail_summary = [
        {"Verdict": "Pass", "Count": data["pass_count"]},
        {"Verdict": "Fail", "Count": data["fail_count"]},
        {"Verdict": "Not computed", "Count": data["not_computed_count"]},
    ]
    sheets = {
        "Header": header,
        "Pass-Fail Summary": pass_fail_summary,
        "Failures by Property": data["property_breakdown"],
        "Failing Results": data["failing_results"],
    }
    charts = [
        ("Pass-Fail Summary", lambda ws, df: _add_bar_chart(ws, df, "Pass / Fail breakdown", "Verdict", ["Count"])),
        ("Failures by Property", lambda ws, df: _add_bar_chart(ws, df, "Failures by property", "Property", ["Fail count"])),
    ]
    if data["show_grade_breakdown"]:
        sheets["Failures by Foam Grade"] = data["grade_breakdown"]
        charts.append((
            "Failures by Foam Grade",
            lambda ws, df: _add_bar_chart(ws, df, "Failures by foam grade", "Foam grade", ["Fail count"]),
        ))
    return _excel_bytes(sheets, charts=charts)


# ---------------------------------------------------------------------------
# 7. PI3 Q&A Report (DOCX only)
# ---------------------------------------------------------------------------

_HTC_LOGO_PATH = "assets/htc_global_logo_blue_steel.png"
_HTC_BLUE = RGBColor(0x1B, 0x6F, 0xA8)  # matches .streamlit/config.toml primaryColor
_HTC_GREY = RGBColor(0x5A, 0x6B, 0x74)


def build_pi3_qa_report_data(
    question, answer, tool_log, page_context="", plant_name=None,
    foam_grade_name=None, asked_by=None, asked_at=None,
):
    """Plain-dict data assembly for one 'Ask PI3' question/answer exchange -
    no Streamlit or python-docx import, so this half is easy to unit test
    on its own. `tool_log` is exactly what ai_assistant.ask_plant_question()
    returns: a list of dicts, each either
    {"tool": "query_plant_data", "sql", "rows_returned", "rows", ["error"]}
    or {"tool": "get_verified_analysis", "args", "result"}."""
    return {
        "question": (question or "").strip(),
        "answer": (answer or "").strip(),
        "tool_log": tool_log or [],
        "page_context": (page_context or "").strip(),
        "plant_name": plant_name,
        "foam_grade_name": foam_grade_name,
        "asked_by": asked_by,
        "asked_at": asked_at or dt.datetime.utcnow(),
    }


def _docx_heading(doc, text, size=13, color=_HTC_BLUE, space_before=12):
    """A styled heading paragraph, built on a real Word "Heading N" style
    (rather than a bold Normal paragraph) so it behaves like an actual
    heading: it shows up in Word's Navigation Pane/outline view and in any
    auto-generated table of contents, and - the practical reason this
    matters here - `keep_with_next` below stops Word from ever stranding a
    heading alone at the bottom of a page with its content pushed to the
    next one.

    Appearance is still fully controlled here via explicit run-level
    formatting (size/color/bold), same as before, so the look stays
    identical regardless of what Word template the opening machine has -
    using a named style doesn't reintroduce that risk, it just makes the
    style semantically real on top of the same explicit formatting."""
    level = 2 if size >= 15 else (3 if size >= 12 else 4)
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


# Recognizes the fixed structure PI3's system prompt always produces (see
# ai_assistant.py's SYSTEM_PROMPT, section "9) Default Response Structure"):
# short, punctuation-free numbered top-level section titles ("1. Direct
# Answer"), optional single-letter-lettered sub-sections ("A. Reduced
# silicone performance"), and "- " prefixed list items. Matching on these
# turns PI3's plain text into real headings/bullets instead of one flat
# Normal paragraph per line - deliberately conservative (short line, starts
# with a capital letter, no internal period, no trailing period) so an
# ordinary sentence that happens to start with a number or single letter
# doesn't get misread as a heading.
_TOP_HEADING_RE = re.compile(r"^\d{1,2}\.\s+[A-Z][^.\n]{2,78}$")
_SUB_HEADING_RE = re.compile(r"^[A-J]\.\s+[A-Z][^.\n]{2,98}$")
_BULLET_RE = re.compile(r"^[-•*]\s+(\S.*)$")

# PI3's free-form question-answering prompt (ai_assistant.PLANT_QUERY_SYSTEM_PROMPT,
# used by the Ask PI3 box) doesn't forbid markdown the way the fixed-prompt
# one does, so its answers commonly contain **bold** / *italic* / `code`
# inline. Without this, those markers came through as literal asterisks/
# backticks in the Word doc instead of real formatting.
_INLINE_MD_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`")

# PI3 also frequently answers with a GFM-style markdown table (a header
# row, a "|---|---:|" alignment/separator row, then data rows) when
# comparing several components or properties - confirmed in production
# output (e.g. a Recipe Optimization report's "Recommended formulation
# direction" and "Target-property focus" tables). Before this existed,
# every line of a markdown table fell through to the plain-paragraph case
# below and rendered as literal "| Water | 3.00 php | ... |" / "|---|---:|"
# text - unreadable, and the single biggest formatting complaint on these
# reports. _PIPE_ROW_RE spots a candidate row; _SEP_CELL_RE recognizes the
# separator row's cells (dashes, optionally colon-flanked for alignment)
# so a real header+data table can be told apart from an ordinary line that
# merely happens to contain a "|" character.
_PIPE_ROW_RE = re.compile(r"^\|.*\|$")
_SEP_CELL_RE = re.compile(r"^:?-{3,}:?$")


def _split_md_table_row(line):
    """Split a markdown table row ('| a | b |') into ['a', 'b'], honoring a
    backslash-escaped pipe inside a cell and dropping the empty strings
    produced by the row's own leading/trailing pipes."""
    protected = line.replace("\\|", "\x00")
    cells = [c.strip().replace("\x00", "|") for c in protected.split("|")]
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return cells


def _is_md_table_separator(line):
    if not _PIPE_ROW_RE.match(line):
        return False
    cells = _split_md_table_row(line)
    return bool(cells) and all(_SEP_CELL_RE.match(c) for c in cells)


def _strip_inline_markdown(text):
    """Plain-text version of a line with markdown markers removed but their
    content kept - used for heading lines, which are already bold/colored
    by their own style, so there's no run-formatting reason to parse
    **bold**/*italic* there, just to avoid showing the literal markers."""
    return _INLINE_MD_RE.sub(lambda m: next(g for g in m.groups() if g is not None), text)


def _add_runs_with_inline_markdown(paragraph, text, size=None):
    """Append `text` to `paragraph` as one or more runs, converting
    **bold**, *italic*, and `code` markdown spans into real run formatting
    instead of leaving the literal markers in the output."""
    pos = 0
    for m in _INLINE_MD_RE.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            if size:
                run.font.size = size
        if m.group(1) is not None:
            run = paragraph.add_run(m.group(1))
            run.bold = True
        elif m.group(2) is not None:
            run = paragraph.add_run(m.group(2))
            run.italic = True
        else:
            run = paragraph.add_run(m.group(3))
            run.font.name = "Consolas"
        if size:
            run.font.size = size
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if size:
            run.font.size = size


def _docx_markdown_table(doc, header_cells, data_rows):
    """Renders a parsed markdown table (a header cell list plus a list of
    data-row cell lists, both already produced by _split_md_table_row) as a
    real bordered Word table - same "Light Grid Accent 1" style used for
    PI3's own SQL-result appendix table (_docx_data_table), so a markdown
    table and a data table look consistent in the same report. Ragged rows
    (a data row with a different cell count than the header) are padded or
    truncated to the header's column count rather than raising - a
    model-written table is exactly the kind of input that occasionally
    comes out uneven."""
    ncols = len(header_cells)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Light Grid Accent 1"
    for cell, header_text in zip(table.rows[0].cells, header_cells):
        run = cell.paragraphs[0].add_run(_strip_inline_markdown(header_text))
        run.bold = True
        run.font.size = Pt(9)
    for row_cells in data_rows:
        padded = (row_cells + [""] * ncols)[:ncols]
        cells = table.add_row().cells
        for cell, cell_text in zip(cells, padded):
            _add_runs_with_inline_markdown(cell.paragraphs[0], cell_text or "—", size=Pt(9))
    return table


def _render_ai_answer_body(doc, text):
    """Render a PI3 answer's plain text into real Word structure: numbered
    top-level sections and lettered sub-sections become real headings,
    "- " list items become real bulleted paragraphs, a markdown pipe table
    (header row + "|---|" separator row + data rows) becomes a real Word
    table, inline **bold**/*italic*/`code` markdown becomes real run
    formatting, and everything else stays a normal paragraph. Replaces the
    previous behavior of dumping one flat Normal paragraph per non-blank
    line verbatim, which produced an unreadable wall of text with no
    headings, bullets, or tables, and left literal markdown markers
    (including whole "| a | b |" / "|---|---:|" table rows) in place."""
    lines = [raw_line.strip() for raw_line in (text or "").split("\n")]
    n = len(lines)
    i = 0
    while i < n:
        line = lines[i]
        if not line:
            i += 1
            continue

        if _PIPE_ROW_RE.match(line) and i + 1 < n and _is_md_table_separator(lines[i + 1]):
            header_cells = _split_md_table_row(line)
            j = i + 2
            data_rows = []
            while j < n and _PIPE_ROW_RE.match(lines[j]):
                data_rows.append(_split_md_table_row(lines[j]))
                j += 1
            _docx_markdown_table(doc, header_cells, data_rows)
            i = j
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_inline_markdown(p, bullet_match.group(1), size=Pt(10.5))
            i += 1
            continue
        if _TOP_HEADING_RE.match(line):
            _docx_heading(doc, _strip_inline_markdown(line), size=13, space_before=14)
            i += 1
            continue
        if _SUB_HEADING_RE.match(line):
            _docx_heading(doc, _strip_inline_markdown(line), size=11.5, color=_HTC_GREY, space_before=10)
            i += 1
            continue
        p = doc.add_paragraph()
        _add_runs_with_inline_markdown(p, line)
        i += 1


def _docx_kv_table(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for label, value in pairs:
        row = table.add_row().cells
        label_run = row[0].paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9.5)
        row[1].paragraphs[0].add_run("—" if value in (None, "") else str(value)).font.size = Pt(9.5)
    return table


def _docx_data_table(doc, rows, max_rows=200):
    """Renders a list-of-dicts as a bordered table, capped at max_rows so a
    very large query result doesn't produce an unusable multi-hundred-page
    appendix - the SQL that produced it is always shown alongside, so the
    full result set is still reproducible."""
    if not rows:
        doc.add_paragraph("No rows returned.").runs[0].font.size = Pt(9)
        return
    shown = rows[:max_rows]
    headers = list(shown[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, header in zip(table.rows[0].cells, headers):
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(8.5)
    for row_data in shown:
        cells = table.add_row().cells
        for cell, header in zip(cells, headers):
            value = row_data.get(header)
            cell.paragraphs[0].add_run("—" if value in (None, "") else str(value)).font.size = Pt(8.5)
    if len(rows) > max_rows:
        note = doc.add_paragraph(f"... {len(rows) - max_rows} further row(s) not shown.")
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(8.5)


def render_pi3_qa_report_docx(data):
    """Renders one PI3 Q&A exchange as DOCX bytes: HTC-branded header,
    a metadata block, the question, PI3's answer, an advisory-boundary
    disclaimer, and an appendix showing exactly what PI3 checked to
    produce the answer. Same layout, same styling, every single time -
    that consistency is the whole point of generating this from code
    rather than starting from a hand-edited Word file each time."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # --- Header: logo + title -------------------------------------------
    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Cm(3.6)
    logo_cell, title_cell = header.rows[0].cells
    if os.path.exists(_HTC_LOGO_PATH):
        run = logo_cell.paragraphs[0].add_run()
        run.add_picture(_HTC_LOGO_PATH, width=Cm(3.0))
    title_p = title_cell.paragraphs[0]
    title_run = title_p.add_run("PI3 Q&A Report")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = _HTC_BLUE
    subtitle_p = title_cell.add_paragraph()
    subtitle_run = subtitle_p.add_run("Flexible slabstock foam expert system | HTC Global Co. Ltd")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = _HTC_GREY

    doc.add_paragraph()

    # --- Metadata ----------------------------------------------------------
    _docx_kv_table(doc, [
        ("Generated", data["asked_at"].strftime("%Y-%m-%d %H:%M UTC")),
        ("Plant", data.get("plant_name") or "—"),
        ("Foam grade", data.get("foam_grade_name") or "—"),
        ("Asked by", data.get("asked_by") or "—"),
        ("Page context", data.get("page_context") or "—"),
    ])

    # --- Question / answer ---------------------------------------------
    # "Question asked"/"PI3's answer"/"Appendix" are the report's top-level
    # sections, so all three sit at Heading 2 - that leaves Heading 3 free
    # for PI3's own numbered sections ("1. Direct Answer") to nest properly
    # underneath "PI3's answer" instead of sitting as its siblings.
    _docx_heading(doc, "Question asked", size=15)
    doc.add_paragraph(data["question"] or "—")

    _docx_heading(doc, "PI3's answer", size=15)
    _render_ai_answer_body(doc, data["answer"] or "—")

    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.space_before = Pt(10)
    disc_run = disclaimer.add_run(
        "This is historical reference for the reviewer's own investigation, not an "
        "instruction. Confirm through your own investigation before acting on it."
    )
    disc_run.italic = True
    disc_run.font.size = Pt(9)
    disc_run.font.color.rgb = _HTC_GREY

    # --- Appendix: exactly what PI3 checked ------------------------------
    doc.add_page_break()
    _docx_heading(doc, "Appendix: data PI3 checked", size=15)
    tool_log = data.get("tool_log") or []
    if not tool_log:
        doc.add_paragraph("No tool calls were recorded for this answer.")
    for i, entry in enumerate(tool_log, start=1):
        tool_name = entry.get("tool", "unknown tool")
        _docx_heading(doc, f"{i}. {tool_name}", size=13, color=_HTC_GREY, space_before=14)
        if tool_name == "query_plant_data":
            sql_p = doc.add_paragraph()
            sql_run = sql_p.add_run(entry.get("sql", "—"))
            sql_run.font.name = "Consolas"
            sql_run.font.size = Pt(9)
            if "error" in entry:
                err_p = doc.add_paragraph()
                err_run = err_p.add_run(f"Rejected: {entry['error']}")
                err_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
                err_run.font.size = Pt(9)
            else:
                count_p = doc.add_paragraph()
                count_run = count_p.add_run(f"{entry.get('rows_returned', 0)} row(s) returned:")
                count_run.font.size = Pt(9)
                _docx_data_table(doc, entry.get("rows") or [])
        elif tool_name == "get_verified_analysis":
            args_p = doc.add_paragraph()
            args_run = args_p.add_run(f"Arguments: {entry.get('args')}")
            args_run.font.size = Pt(9)
            result = entry.get("result")
            if isinstance(result, dict):
                kv_pairs = []
                table_keys = []  # list-of-dicts values, rendered as a sub-table below instead
                for k, v in result.items():
                    if isinstance(v, list) and v and isinstance(v[0], dict):
                        kv_pairs.append((k, f"[{len(v)} row(s) - see table below]"))
                        table_keys.append(k)
                    elif isinstance(v, list):
                        # Plain-value list (e.g. warnings/successes strings) - show the
                        # actual content inline rather than hiding it behind a count.
                        kv_pairs.append((k, "; ".join(str(x) for x in v) if v else "—"))
                    else:
                        kv_pairs.append((k, v))
                _docx_kv_table(doc, kv_pairs)
                for k in table_keys:
                    sub = doc.add_paragraph()
                    sub_run = sub.add_run(k)
                    sub_run.bold = True
                    sub_run.font.size = Pt(9)
                    _docx_data_table(doc, result[k])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
